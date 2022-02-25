'''
Auther: Haorong Jiang
Date: 2022-02-22 21:40:48
LastEditors: Haorong Jiang
LastEditTime: 2022-02-24 21:45:46
'''
import os
import re
import shutil
import sys
from tkinter.tix import TEXT
from xml.dom.minidom import Document

import docx
import pygame
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH



class Color:
    # 自定义颜色
    ACHIEVEMENT = (220, 160, 87)
    VERSION = (220, 160, 87)

    # 固定颜色
    BLACK = (0, 0, 0)
    WHITE = (255, 255, 255)
    RED = (255, 0, 0)
    GREEN = (0, 255, 0)
    BLUE = (0, 0, 255)
    GREY = (128, 128, 128)  # 中性灰
    TRANSPARENT = (255, 255, 255, 0)  # 白色的完全透明


class Text:
    def __init__(self, text: str, text_color: Color, font_type: str, font_size: int):
        """
        text: 文本内容，如'大学生模拟器'，注意是字符串形式
        text_color: 字体颜色，如Color.WHITE、COLOR.BLACK
        font_type: 字体文件(.ttc)，如'msyh.ttc'，注意是字符串形式
        font_size: 字体大小，如20、10
        """
        self.text = text
        self.text_color = text_color
        self.font_type = font_type
        self.font_size = font_size

        font = pygame.font.Font(os.path.join('font', (self.font_type)), self.font_size)
        self.text_image = font.render(self.text, True, self.text_color).convert_alpha()

        self.text_width = self.text_image.get_width()
        self.text_height = self.text_image.get_height()

    def draw(self, surface: pygame.Surface, center_x, center_y):
        """
        surface: 文本放置的表面
        center_x, center_y: 文本放置在表面的<中心坐标>
        """
        upperleft_x = center_x - self.text_width / 2
        upperleft_y = center_y - self.text_height / 2
        surface.blit(self.text_image, (upperleft_x, upperleft_y))
    
    def upperleft_draw(self, surface: pygame.Surface, upperleft_x, upperleft_y):
        surface.blit(self.text_image, (upperleft_x, upperleft_y))


class ColorSurface:
    def __init__(self, color, width, height):
        self.color = color
        self.width = width
        self.height = height

        self.color_image = pygame.Surface((self.width, self.height)).convert_alpha()
        self.color_image.fill(self.color)

    def draw(self, surface: pygame.Surface, center_x, center_y):
        upperleft_x = center_x - self.width / 2
        upperleft_y = center_y - self.height / 2
        surface.blit(self.color_image, (upperleft_x, upperleft_y))


class ButtonColorSurface(ColorSurface):
    def __init__(self, color, width, height):
        super().__init__(color, width, height)
        self.rect = self.color_image.get_rect()

    def draw(self, surface: pygame.Surface, center_x, center_y):
        super().draw(surface, center_x, center_y)
        self.rect.center = center_x, center_y

    def handle_event(self, command, *args):
        self.hovered = self.rect.collidepoint(pygame.mouse.get_pos())
        if self.hovered:
            command(*args)


def docx_process(screen, center_x, center_y): 
    source_dir = os.path.join('docx_files')
    result_dir = os.path.join('result')

    print(source_dir)

    # 每一次操作前，清空result文件夹
    shutil.rmtree(result_dir)  
    os.mkdir(result_dir)  

    num = 0

    for root, dirs, files in os.walk(source_dir):
        for file in files:
            if file.endswith('.docx'):
                
                num += 1

                current_source_docx = os.path.join(root, file)
                current_result_docx = os.path.join(result_dir, file)

                doc = docx.Document(
                    current_source_docx
                )

                doc.styles['Normal'].font.name = u'Times New Roman'

                for para in doc.paragraphs:
                    para.paragraph_format.line_spacing = 1.0

                # kill all parentheses
                for para in doc.paragraphs:
                    para_no_parenthese = re.sub(u"\\(.*?\\)|\\{.*?}|\\[.*?]", "", para.text)
                    para.text = para_no_parenthese

                # kill references and after
                ref_flag = False
                for para in doc.paragraphs:
                    if str(para.text).lower() == 'references':
                        ref_flag = True
                    if ref_flag == True:
                        para.clear()

                # kill before abstract
                for para in doc.paragraphs:
                    if str(para.text).lower() == 'abstract':
                        break
                    else:
                        para.clear()

                # kill all paragraphs starts with number & table paragraphs
                for para in doc.paragraphs:
                    if len(para.text) > 0:
                        if str(para.text)[0] in [
                                '0', '1', '2', '3', '4', '5', '6', '7', '8', '9'
                        ]:
                            para.clear()
                        if str(para.text)[:5].lower() == 'table':
                            para.clear()
                        if str(para.text)[:6].lower() == 'figure':
                            para.clear()

                # kill all tables
                for table in doc.tables:
                    table._element.getparent().remove(table._element)

                # for para in doc.paragraphs:
                #     print(para.text)

                doc.save(current_result_docx)
                
                ColorSurface((210, 210, 210), 200, 30).draw(screen, center_x, center_y)
                Text('目前处理: '+ str(num), Color.BLACK, 'HYHanHeiW.ttf', 20).draw(screen, center_x, center_y)


class InterFace():
    def __init__(self):
        pygame.init()

    def basic_interface(self):
        pygame.display.set_caption('docx文件预处理')
        size = width, height = 600, 371
        screen = pygame.display.set_mode(size)

        ColorSurface((210, 210, 210), width, height).draw(screen, width * 0.5, height * 0.5)  # 背景颜色
        Text('docx文件预处理', Color.BLACK, 'HYHanHeiW.ttf', 40).draw(screen, width * 0.3, height * 0.15) # 标题颜色
        Text('V1.0', Color.BLACK, 'HYHanHeiW.ttf', 15).draw(screen, width * 0.6, height * 0.15)

        return size, screen

    def start_interface(self):
        size, screen = self.basic_interface()
        width, height = size

        # 版权申明
        Text('By Haorong Jiang', Color.BLACK, 'HYHanHeiW.ttf', 13).draw(screen, width * 0.86, height * 0.85)
        Text('All Rights Reserved.', Color.BLACK, 'HYHanHeiW.ttf', 13).draw(screen, width * 0.86, height * 0.9)
        
        # 使用说明按钮
        shiyongshuoming = ButtonColorSurface((150, 150, 150), 160, 40)
        shiyongshuoming.draw(screen, width * 0.3, height * 0.45)
        Text('使用说明', Color.BLACK, 'HYHanHeiW.ttf', 25).draw(screen, width * 0.3, height * 0.45)

        # 开始处理按钮
        kaishichuli = ButtonColorSurface((23, 186, 1), 160, 40)
        kaishichuli.draw(screen, width * 0.7, height * 0.45)
        Text('开始处理', Color.BLACK, 'HYHanHeiW.ttf', 25).draw(screen, width * 0.7, height * 0.45)

        while True:
            for event in pygame.event.get():
                if event.type == pygame.QUIT:
                    pygame.quit()
                    sys.exit()

                if event.type == pygame.MOUSEBUTTONDOWN:
                    shiyongshuoming.handle_event(self.shiyongshuoming_interface)
                    kaishichuli.handle_event(self.kaishichuli_interface)

            pygame.display.update()

    def kaishichuli_interface(self):
        size, screen = self.basic_interface()
        width, height = size

        # 版权申明
        Text('By Haorong Jiang', Color.BLACK, 'HYHanHeiW.ttf', 13).draw(screen, width * 0.86, height * 0.85)
        Text('All rights reserved.', Color.BLACK, 'HYHanHeiW.ttf', 13).draw(screen, width * 0.86, height * 0.9)
        
        # 使用说明按钮
        shiyongshuoming = ButtonColorSurface((150, 150, 150), 160, 40)
        shiyongshuoming.draw(screen, width * 0.3, height * 0.45)
        Text('使用说明', Color.BLACK, 'HYHanHeiW.ttf', 25).draw(screen, width * 0.3, height * 0.45)

        # 开始处理按钮
        kaishichuli = ButtonColorSurface((23, 186, 1), 160, 40)
        kaishichuli.draw(screen, width * 0.7, height * 0.45)
        Text('开始处理', Color.BLACK, 'HYHanHeiW.ttf', 25).draw(screen, width * 0.7, height * 0.45)


        Text('准备处理...', Color.BLACK, 'HYHanHeiW.ttf', 20).draw(screen, width * 0.5, height * 0.6)

        docx_process(screen, width * 0.5, height * 0.7)

        Text('处理完成', Color.BLACK, 'HYHanHeiW.ttf', 20).draw(screen, width * 0.5, height * 0.8)
        
        
        
    def shiyongshuoming_interface(self):
        size, screen = self.basic_interface()
        width, height = size

        # 返回按钮
        fanhui = ButtonColorSurface((150, 150, 150), 160, 40)
        fanhui.draw(screen, width * 0.8, height * 0.3)
        Text('返回主界面', Color.BLACK, 'HYHanHeiW.ttf', 25).draw(screen, width * 0.8, height * 0.3)

        # 说明文字
        Text('使用说明', Color.BLACK, 'HYHanHeiW.ttf', 25).upperleft_draw(screen, width * 0.1, height * 0.26)
        Text('1. 请不要改变任何文件夹名称和结构', Color.BLACK, 'HYHanHeiW.ttf', 18).upperleft_draw(screen, width * 0.1, height * 0.38)
        Text('2. 将想要预处理的docx文件（只能是docx文件）', Color.BLACK, 'HYHanHeiW.ttf', 18).upperleft_draw(screen, width * 0.1, height * 0.46)
        Text('    放于<docx_files>文件夹中', Color.BLACK, 'HYHanHeiW.ttf', 18).upperleft_draw(screen, width * 0.1, height * 0.54)
        Text('3. 点击开始<开始处理>按钮', Color.BLACK, 'HYHanHeiW.ttf', 18).upperleft_draw(screen, width * 0.1, height * 0.62)
        Text('4. 显示<处理完成>后，对应的docx文件', Color.BLACK, 'HYHanHeiW.ttf', 18).upperleft_draw(screen, width * 0.1, height * 0.7)
        Text('    生成在<result>文件夹中', Color.BLACK, 'HYHanHeiW.ttf', 18).upperleft_draw(screen, width * 0.1, height * 0.78)

        while True:
            for event in pygame.event.get():
                if event.type == pygame.QUIT:
                    pygame.quit()
                    sys.exit()

                if event.type == pygame.MOUSEBUTTONDOWN:
                    fanhui.handle_event(self.start_interface)

            pygame.display.update()


if __name__ == '__main__':
    scene = InterFace()
    scene.start_interface()
